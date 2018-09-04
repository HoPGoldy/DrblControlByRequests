# -*- coding: utf-8 -*-
# @Time    : 2018/9/3 12:45
# @Author  : HoPGoldy

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

DETAIL_TEMPLATE = '''单品表的word文档导入
[标题]：{title}
[自定义分类]：{category}
[本文目标人群]: {target1} >> {target2}
[商品宝贝]:
{url}
[好在哪里]: 
{LHL}
[短亮点]: {SHL}
[段落]: '''

ADDITION_TEMPLATE = '''【段落{additionCode}】：
<段落标题>: {additionTitle}
<段落描述>: 
{additionContent}
<段落随图>:
{additionImg}
'''

CODE_ARRAY = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']


def createWord(data, wordName):
    detailTemp = DETAIL_TEMPLATE
    detailTemp = detailTemp.replace('{title}', data['title'])
    detailTemp = detailTemp.replace('{category}', data['category'])
    detailTemp = detailTemp.replace('{url}', data['url'])
    detailTemp = detailTemp.replace('{target1}', data['targetPeople'][0])
    detailTemp = detailTemp.replace('{target2}', data['targetPeople'][1])
    LHL = ''
    for longHighLight in data['longHighLight']:
        LHL += longHighLight + r'//'
    detailTemp = detailTemp.replace('{LHL}', LHL)
    SHL = ''
    for shortHighLight in data['longHighLight']:
        SHL += shortHighLight + r'//'
    detailTemp = detailTemp.replace('{SHL}', SHL)

    additionTemp = []
    for i in range(len(data['addition'])):
        addition = data['addition'][i]
        temp = ADDITION_TEMPLATE
        temp = temp.replace('{additionCode}', CODE_ARRAY[i])
        temp = temp.replace('{additionTitle}', addition['title'])
        temp = temp.replace('{additionContent}', addition['content'])
        temp = temp.replace('{additionImg}', addition['img'])
        additionTemp.append(temp)

    # 打开文档
    document = Document()
    # 设置字体字号
    document.styles['Normal'].font.name = u'黑体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    document.styles['Normal'].font.size = Pt(14)

    # 添加内容
    for paragraph in detailTemp.split('\n'):
        document.add_paragraph(paragraph)
    for addition in additionTemp:
        for paragraph in addition.split('\n'):
            document.add_paragraph(paragraph)
            # 增加图像
            # document.add_picture('1.jpg')

    # 保存文件
    document.save(wordName)